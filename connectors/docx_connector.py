"""
E-T-A RAG Prototype - Word Document Connector

From Architecture Design (Section 5.1):
- Extraction: python-docx / pandoc
- Chunking: Section-aware splitting, 512-token chunks
- Output: Structured text
"""
import hashlib
import os
from pathlib import Path
from typing import Optional

from connectors.base import BaseConnector, DocumentChunk


class DocxConnector(BaseConnector):
    """Connector for Word (.docx) files."""

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def extract(self, file_path: str, metadata: Optional[dict] = None) -> list[DocumentChunk]:
        from docx import Document

        doc = Document(file_path)
        filename = os.path.basename(file_path)
        chunks = []

        # Extract text by sections/headings
        current_section = ""
        current_text = ""
        section_idx = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if this is a heading (new section)
            if para.style.name.startswith("Heading"):
                # Save previous section
                if current_text.strip():
                    chunks.extend(
                        self._text_to_chunks(
                            current_text, filename, current_section, section_idx, metadata
                        )
                    )
                    section_idx += 1
                current_section = text
                current_text = f"## {text}\n\n"
            else:
                current_text += text + "\n\n"

        # Don't forget last section
        if current_text.strip():
            chunks.extend(
                self._text_to_chunks(
                    current_text, filename, current_section, section_idx, metadata
                )
            )

        # Also extract tables
        for table_idx, table in enumerate(doc.tables):
            table_text = self._extract_table(table)
            if table_text:
                chunk_id = hashlib.md5(
                    f"{file_path}:table:{table_idx}".encode()
                ).hexdigest()
                chunk = DocumentChunk(
                    text=f"[Source: {filename}, Table {table_idx + 1}]\n{table_text}",
                    source_format="docx",
                    source_file=filename,
                    page_number=table_idx,
                    chunk_id=chunk_id,
                )
                self._apply_metadata(chunk, metadata)
                chunks.append(chunk)

        return chunks

    def _text_to_chunks(
        self, text: str, filename: str, section: str, section_idx: int,
        metadata: Optional[dict] = None
    ) -> list[DocumentChunk]:
        """Split section text into chunks."""
        max_chars = 2048
        parts = [text] if len(text) <= max_chars else self._split(text, max_chars)
        chunks = []

        for i, part in enumerate(parts):
            chunk_id = hashlib.md5(
                f"{filename}:s{section_idx}:c{i}".encode()
            ).hexdigest()
            chunk = DocumentChunk(
                text=f"[Source: {filename}, Section: {section}]\n{part}",
                source_format="docx",
                source_file=filename,
                page_number=section_idx,
                chunk_id=chunk_id,
            )
            self._apply_metadata(chunk, metadata)
            chunks.append(chunk)
        return chunks

    def _split(self, text: str, max_chars: int) -> list[str]:
        paragraphs = text.split("\n\n")
        result, current = [], ""
        for p in paragraphs:
            if len(current) + len(p) > max_chars and current:
                result.append(current.strip())
                current = p
            else:
                current = f"{current}\n\n{p}" if current else p
        if current.strip():
            result.append(current.strip())
        return result

    def _extract_table(self, table) -> str:
        """Extract table content as readable text."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows) if rows else ""
